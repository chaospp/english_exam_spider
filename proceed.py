from docx import Document
from docx.shared import Pt,RGBColor

def proceed():
    document = Document()  # 新建文档对象
    document.add_heading('测验', level=4) # 设置标题级别
    f = open('0.txt','r', encoding="utf-8")
    f2 = open('1.txt','r', encoding="utf-8")
    p1 = document.add_paragraph(f)
    p2 = document.add_paragraph(f2)
    format2 = p2.paragraph_format
    # 左右缩进
    format2.left_indent = Pt(20)
    format2.right_indent = Pt(20)
    format2.first_line_indent = Pt(20) # 首行缩进
    # 行间距
    format2.line_spacing = 1.5
    format = p1.paragraph_format
    # 左右缩进
    format.left_indent = Pt(20)
    format.right_indent = Pt(20)
    format.first_line_indent = Pt(20) # 首行缩进
    # 行间距
    format.line_spacing = 1.5
    # 3.保存文档
    document.save('info.docx')
    print('保存为info.dox')
